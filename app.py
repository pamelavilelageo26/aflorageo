import streamlit as st
import sqlite3
import os
import json
import datetime
import base64
import io
from pathlib import Path

# --- CONFIGURAÇÃO INICIAL ---
st.set_page_config(
    page_title="AfloraGeo - Caderneta de Campo Geológica",
    page_icon="🌋",
    layout="wide",
    initial_sidebar_state="expanded",
)

DB_FILE = "aflorageo.db"

# --- CSS PERSONALIZADO ---
st.markdown(
    """
    <style>
    .stButton>button { width: 100%; }
    div[data-testid="stMetric"] { background-color: #f0f2f6; padding: 10px; border-radius: 10px; }
    .premium-card { border: 2px solid #e0e0e0; border-radius: 12px; padding: 16px; text-align: center; }
    .premium-card-popular { border: 2px solid #ffd700; background-color: #fffdf5; border-radius: 12px; padding: 16px; text-align: center; }
    </style>
    """,
    unsafe_allow_html=True,
)

# --- FUNÇÕES DE BANCO DE DADOS ---

def get_connection():
    return sqlite3.connect(DB_FILE, check_same_thread=False)

def init_db():
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute('''
            CREATE TABLE IF NOT EXISTS stations (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                ponto_id TEXT UNIQUE,
                data TEXT,
                utm_zone TEXT,
                hemisferio TEXT,
                utm_east REAL,
                utm_north REAL,
                latitude REAL,
                longitude REAL,
                altitude REAL,
                localizacao TEXT,
                municipio TEXT,
                contexto_geologico TEXT,
                tipo_afloramento TEXT,
                dimensoes TEXT,
                orientacao_afloramento TEXT,
                acesso TEXT,
                litologia_principal TEXT,
                litologia_secundaria TEXT,
                granulometria TEXT,
                cor TEXT,
                intemperismo TEXT,
                observacoes TEXT,
                created_at TEXT,
                updated_at TEXT
            )
        ''')
        cur.execute('''
            CREATE TABLE IF NOT EXISTS structures (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                station_id INTEGER,
                tipo TEXT,
                strike TEXT,
                dip REAL,
                dip_dir REAL,
                plunge REAL,
                azimuth REAL,
                observacoes TEXT,
                created_at TEXT
            )
        ''')
        cur.execute('''
            CREATE TABLE IF NOT EXISTS samples (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                station_id INTEGER,
                codigo TEXT,
                tipo TEXT,
                finalidade TEXT,
                orientada INTEGER,
                observacoes TEXT,
                created_at TEXT
            )
        ''')
        cur.execute('''
            CREATE TABLE IF NOT EXISTS photos (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                station_id INTEGER,
                descricao TEXT,
                arquivo TEXT,
                created_at TEXT
            )
        ''')
        cur.execute('''
            CREATE TABLE IF NOT EXISTS license (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                is_premium INTEGER DEFAULT 0,
                stations_limit INTEGER DEFAULT 30,
                user_email TEXT,
                plan_type TEXT,
                expires_at TEXT,
                created_at TEXT,
                updated_at TEXT
            )
        ''')
        conn.commit()
    except Exception as e:
        st.error(f"Erro ao inicializar banco de dados: {e}")
    finally:
        if conn:
            conn.close()

def init_license():
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("SELECT COUNT(*) FROM license")
        if cur.fetchone()[0] == 0:
            now = datetime.datetime.now().isoformat()
            cur.execute('''
                INSERT INTO license (is_premium, stations_limit, user_email, plan_type, created_at, updated_at)
                VALUES (0, 30, '', 'free', ?, ?)
            ''', (now, now))
            conn.commit()
    except Exception as e:
        st.error(f"Erro ao inicializar licença: {e}")
    finally:
        if conn:
            conn.close()

def get_license():
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("SELECT * FROM license ORDER BY id DESC LIMIT 1")
        row = cur.fetchone()
        if row:
            return {
                "id": row[0],
                "is_premium": bool(row[1]),
                "stations_limit": row[2],
                "user_email": row[3],
                "plan_type": row[4],
                "expires_at": row[5],
                "created_at": row[6],
                "updated_at": row[7],
            }
        return None
    except Exception as e:
        st.error(f"Erro ao consultar licença: {e}")
        return None
    finally:
        if conn:
            conn.close()

def is_premium():
    lic = get_license()
    if lic and lic["is_premium"]:
        return True
    return False

def get_station_count():
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("SELECT COUNT(*) FROM stations")
        return cur.fetchone()[0]
    except Exception as e:
        st.error(f"Erro ao contar estações: {e}")
        return 0
    finally:
        if conn:
            conn.close()

def can_add_station():
    lic = get_license()
    if lic and lic["is_premium"]:
        return True
    count = get_station_count()
    limit = lic["stations_limit"] if lic else 30
    return count < limit

def activate_premium(plan_type, email):
    try:
        conn = get_connection()
        cur = conn.cursor()
        now = datetime.datetime.now()
        if plan_type == "monthly":
            expires = now + datetime.timedelta(days=30)
        elif plan_type == "semiannual":
            expires = now + datetime.timedelta(days=180)
        elif plan_type == "annual":
            expires = now + datetime.timedelta(days=365)
        else:
            expires = now + datetime.timedelta(days=30)
        cur.execute("DELETE FROM license")
        cur.execute('''
            INSERT INTO license (is_premium, stations_limit, user_email, plan_type, expires_at, created_at, updated_at)
            VALUES (1, 999999, ?, ?, ?, ?, ?)
        ''', (email, plan_type, expires.isoformat(), now.isoformat(), now.isoformat()))
        conn.commit()
        return True
    except Exception as e:
        st.error(f"Erro ao ativar premium: {e}")
        return False
    finally:
        if conn:
            conn.close()

def reset_license():
    try:
        conn = get_connection()
        cur = conn.cursor()
        now = datetime.datetime.now().isoformat()
        cur.execute("DELETE FROM license")
        cur.execute('''
            INSERT INTO license (is_premium, stations_limit, user_email, plan_type, created_at, updated_at)
            VALUES (0, 30, '', 'free', ?, ?)
        ''', (now, now))
        conn.commit()
    except Exception as e:
        st.error(f"Erro ao resetar licença: {e}")
    finally:
        if conn:
            conn.close()

def next_station_id(cur):
    cur.execute("SELECT MAX(id) FROM stations")
    max_id = cur.fetchone()[0]
    if max_id is None:
        next_num = 1
    else:
        next_num = max_id + 1
    return f"AF-{next_num:03d}"

def list_stations_df():
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("SELECT * FROM stations ORDER BY id DESC")
        rows = cur.fetchall()
        columns = [desc[0] for desc in cur.description]
        import pandas as pd
        return pd.DataFrame(rows, columns=columns)
    except Exception as e:
        st.error(f"Erro ao listar estações: {e}")
        return None
    finally:
        if conn:
            conn.close()

def get_station_by_id(station_id):
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("SELECT * FROM stations WHERE id = ?", (station_id,))
        row = cur.fetchone()
        if row:
            columns = [desc[0] for desc in cur.description]
            return dict(zip(columns, row))
        return None
    except Exception as e:
        st.error(f"Erro ao buscar estação: {e}")
        return None
    finally:
        if conn:
            conn.close()

def get_structures_by_station(station_id):
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("SELECT * FROM structures WHERE station_id = ?", (station_id,))
        rows = cur.fetchall()
        columns = [desc[0] for desc in cur.description]
        return [dict(zip(columns, row)) for row in rows]
    except Exception as e:
        st.error(f"Erro ao buscar estruturas: {e}")
        return []
    finally:
        if conn:
            conn.close()

def get_samples_by_station(station_id):
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("SELECT * FROM samples WHERE station_id = ?", (station_id,))
        rows = cur.fetchall()
        columns = [desc[0] for desc in cur.description]
        return [dict(zip(columns, row)) for row in rows]
    except Exception as e:
        st.error(f"Erro ao buscar amostras: {e}")
        return []
    finally:
        if conn:
            conn.close()

def insert_station(data, structures, samples_list):
    try:
        conn = get_connection()
        cur = conn.cursor()
        ponto_id = next_station_id(cur)
        now = datetime.datetime.now().isoformat()
        cur.execute('''
            INSERT INTO stations (
                ponto_id, data, utm_zone, hemisferio, utm_east, utm_north, latitude, longitude, altitude,
                localizacao, municipio, contexto_geologico, tipo_afloramento, dimensoes, orientacao_afloramento,
                acesso, litologia_principal, litologia_secundaria, granulometria, cor, intemperismo,
                observacoes, created_at, updated_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            ponto_id, data.get("data"), data.get("utm_zone"), data.get("hemisferio"),
            data.get("utm_east"), data.get("utm_north"), data.get("latitude"), data.get("longitude"),
            data.get("altitude"), data.get("localizacao"), data.get("municipio"), data.get("contexto_geologico"),
            data.get("tipo_afloramento"), data.get("dimensoes"), data.get("orientacao_afloramento"),
            data.get("acesso"), data.get("litologia_principal"), data.get("litologia_secundaria"),
            data.get("granulometria"), data.get("cor"), data.get("intemperismo"), data.get("observacoes"),
            now, now
        ))
        station_id = cur.lastrowid
        for s in structures:
            cur.execute('''
                INSERT INTO structures (station_id, tipo, strike, dip, dip_dir, plunge, azimuth, observacoes, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (station_id, s.get("tipo"), s.get("strike"), s.get("dip"), s.get("dip_dir"),
                  s.get("plunge"), s.get("azimuth"), s.get("observacoes"), now))
        for s in samples_list:
            cur.execute('''
                INSERT INTO samples (station_id, codigo, tipo, finalidade, orientada, observacoes, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (station_id, s.get("codigo"), s.get("tipo"), s.get("finalidade"), s.get("orientada"), s.get("observacoes"), now))
        conn.commit()
        return True
    except Exception as e:
        st.error(f"Erro ao salvar estação: {e}")
        return False
    finally:
        if conn:
            conn.close()

def update_station(station_id, data, structures, samples_list):
    try:
        conn = get_connection()
        cur = conn.cursor()
        now = datetime.datetime.now().isoformat()
        cur.execute('''
            UPDATE stations SET
                data = ?, utm_zone = ?, hemisferio = ?, utm_east = ?, utm_north = ?, latitude = ?, longitude = ?,
                altitude = ?, localizacao = ?, municipio = ?, contexto_geologico = ?, tipo_afloramento = ?,
                dimensoes = ?, orientacao_afloramento = ?, acesso = ?, litologia_principal = ?, litologia_secundaria = ?,
                granulometria = ?, cor = ?, intemperismo = ?, observacoes = ?, updated_at = ?
            WHERE id = ?
        ''', (
            data.get("data"), data.get("utm_zone"), data.get("hemisferio"), data.get("utm_east"),
            data.get("utm_north"), data.get("latitude"), data.get("longitude"), data.get("altitude"),
            data.get("localizacao"), data.get("municipio"), data.get("contexto_geologico"), data.get("tipo_afloramento"),
            data.get("dimensoes"), data.get("orientacao_afloramento"), data.get("acesso"), data.get("litologia_principal"),
            data.get("litologia_secundaria"), data.get("granulometria"), data.get("cor"), data.get("intemperismo"),
            data.get("observacoes"), now, station_id
        ))
        cur.execute("DELETE FROM structures WHERE station_id = ?", (station_id,))
        cur.execute("DELETE FROM samples WHERE station_id = ?", (station_id,))
        for s in structures:
            cur.execute('''
                INSERT INTO structures (station_id, tipo, strike, dip, dip_dir, plunge, azimuth, observacoes, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (station_id, s.get("tipo"), s.get("strike"), s.get("dip"), s.get("dip_dir"),
                  s.get("plunge"), s.get("azimuth"), s.get("observacoes"), now))
        for s in samples_list:
            cur.execute('''
                INSERT INTO samples (station_id, codigo, tipo, finalidade, orientada, observacoes, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (station_id, s.get("codigo"), s.get("tipo"), s.get("finalidade"), s.get("orientada"), s.get("observacoes"), now))
        conn.commit()
        return True
    except Exception as e:
        st.error(f"Erro ao atualizar estação: {e}")
        return False
    finally:
        if conn:
            conn.close()

def delete_station(station_id):
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("DELETE FROM structures WHERE station_id = ?", (station_id,))
        cur.execute("DELETE FROM samples WHERE station_id = ?", (station_id,))
        cur.execute("DELETE FROM photos WHERE station_id = ?", (station_id,))
        cur.execute("DELETE FROM stations WHERE id = ?", (station_id,))
        conn.commit()
        return True
    except Exception as e:
        st.error(f"Erro ao excluir estação: {e}")
        return False
    finally:
        if conn:
            conn.close()

def reset_all_data():
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("DELETE FROM structures")
        cur.execute("DELETE FROM samples")
        cur.execute("DELETE FROM photos")
        cur.execute("DELETE FROM stations")
        conn.commit()
    except Exception as e:
        st.error(f"Erro ao resetar dados: {e}")
    finally:
        if conn:
            conn.close()

# --- FUNÇÕES DE CONVERSÃO UTM ---

def utm_to_dd(zone, easting, northing, hemisphere="S"):
    try:
        import utm
        lat, lon = utm.to_latlon(easting, northing, int(zone[:-1]), zone[-1])
        if hemisphere == "S" and lat > 0:
            lat = -lat
        return lat, lon
    except Exception:
        return None, None

def dd_to_utm(lat, lon):
    try:
        import utm
        easting, northing, zone_number, zone_letter = utm.from_latlon(lat, lon)
        zone = f"{zone_number}{zone_letter}"
        hemisphere = "S" if lat < 0 else "N"
        return zone, hemisphere, easting, northing
    except Exception:
        return None, None, None, None

# --- FUNÇÕES DE EXPORTAÇÃO ---

def generate_kml(stations_df, structures_df, samples_df):
    try:
        import simplekml
        kml = simplekml.Kml()
        for _, row in stations_df.iterrows():
            lat = row.get("latitude")
            lon = row.get("longitude")
            alt = row.get("altitude") or 0
            if lat is None or lon is None:
                continue
            pnt = kml.newpoint(name=str(row.get("ponto_id", "")), coords=[(lon, lat, alt)])
            st_structures = structures_df[structures_df["station_id"] == row["id"]] if structures_df is not None and not structures_df.empty else None
            st_samples = samples_df[samples_df["station_id"] == row["id"]] if samples_df is not None and not samples_df.empty else None
            desc = f"""
            <h3>{row.get('ponto_id', '')}</h3>
            <p><b>Data:</b> {row.get('data', '')}</p>
            <p><b>Localização:</b> {row.get('localizacao', '')}</p>
            <p><b>Município:</b> {row.get('municipio', '')}</p>
            <p><b>Coordenadas:</b> {lat:.6f}, {lon:.6f}</p>
            <p><b>Altitude:</b> {alt}</p>
            <p><b>Litologia principal:</b> {row.get('litologia_principal', '')}</p>
            <p><b>Litologia secundária:</b> {row.get('litologia_secundaria', '')}</p>
            <p><b>Tipo de afloramento:</b> {row.get('tipo_afloramento', '')}</p>
            <p><b>Contexto geológico:</b> {row.get('contexto_geologico', '')}</p>
            <p><b>Observações:</b> {row.get('observacoes', '')}</p>
            """
            if st_structures is not None and not st_structures.empty:
                desc += "<<h4>Estruturas</h4><ul>"
                for _, s in st_structures.iterrows():
                    desc += f"<<li>{s.get('tipo', '')} - Strike: {s.get('strike', '')}, Dip: {s.get('dip', '')}, Dip Dir: {s.get('dip_dir', '')}, Plunge: {s.get('plunge', '')}, Azimuth: {s.get('azimuth', '')}</li>"
                desc += "</ul>"
            if st_samples is not None and not st_samples.empty:
                desc += "<<h4>Amostras</h4><ul>"
                for _, s in st_samples.iterrows():
                    desc += f"<<li>{s.get('codigo', '')} - {s.get('tipo', '')} - {s.get('finalidade', '')} - Orientada: {'Sim' if s.get('orientada') else 'Não'}</li>"
                desc += "</ul>"
            pnt.description = desc
        return kml.kml()
    except Exception:
        lines = [
            '<?xml version="1.0" encoding="UTF-8"?>',
            '<kml xmlns="http://www.opengis.net/kml/2.2">',
            "<<Document>",
            "<<name>AfloraGeo</name>",
        ]
        for _, row in stations_df.iterrows():
            lat = row.get("latitude")
            lon = row.get("longitude")
            alt = row.get("altitude") or 0
            if lat is None or lon is None:
                continue
            name = str(row.get("ponto_id", "")).replace("&", "&amp;").replace("<<", "&lt;").replace(">", "&gt;")
            desc = f"Ponto: {name}<br/>Data: {row.get('data', '')}<br/>Localização: {row.get('localizacao', '')}<br/>Município: {row.get('municipio', '')}<br/>Litologia: {row.get('litologia_principal', '')}"
            desc = desc.replace("&", "&amp;").replace("<<", "&lt;").replace(">", "&gt;")
            lines.append("<<Placemark>")
            lines.append(f"<<name>{name}</name>")
            lines.append(f"<<description><![CDATA[{desc}]]></description>")
            lines.append("<<Point>")
            lines.append(f"<<coordinates>{lon},{lat},{alt}</coordinates>")
            lines.append("</Point>")
            lines.append("</Placemark>")
        lines.append("</Document>")
        lines.append("</kml>")
        return "\n".join(lines)

def generate_geojson(stations_df):
    features = []
    for _, row in stations_df.iterrows():
        lat = row.get("latitude")
        lon = row.get("longitude")
        if lat is None or lon is None:
            continue
        features.append({
            "type": "Feature",
            "geometry": {
                "type": "Point",
                "coordinates": [lon, lat, row.get("altitude") or 0]
            },
            "properties": {
                "ponto_id": row.get("ponto_id", ""),
                "data": row.get("data", ""),
                "localizacao": row.get("localizacao", ""),
                "municipio": row.get("municipio", ""),
                "litologia_principal": row.get("litologia_principal", ""),
                "tipo_afloramento": row.get("tipo_afloramento", ""),
                "contexto_geologico": row.get("contexto_geologico", ""),
            }
        })
    return json.dumps({"type": "FeatureCollection", "features": features}, ensure_ascii=False, indent=2)

def generate_csv(stations_df):
    import pandas as pd
    return stations_df.to_csv(index=False, encoding="utf-8-sig")

def generate_docx(estacao_dict, structures_list, samples_list):
    try:
        from docx import Document
        doc = Document()
        doc.add_heading(f"Estação {estacao_dict.get('ponto_id', '')}", level=1)
        doc.add_paragraph(f"Data: {estacao_dict.get('data', '')}")
        doc.add_paragraph(f"Localização: {estacao_dict.get('localizacao', '')}")
        doc.add_paragraph(f"Município: {estacao_dict.get('municipio', '')}")
        doc.add_paragraph(f"Coordenadas: {estacao_dict.get('latitude', '')}, {estacao_dict.get('longitude', '')}")
        doc.add_paragraph(f"Altitude: {estacao_dict.get('altitude', '')}")
        doc.add_paragraph(f"Litologia principal: {estacao_dict.get('litologia_principal', '')}")
        doc.add_paragraph(f"Litologia secundária: {estacao_dict.get('litologia_secundaria', '')}")
        doc.add_paragraph(f"Tipo de afloramento: {estacao_dict.get('tipo_afloramento', '')}")
        doc.add_paragraph(f"Contexto geológico: {estacao_dict.get('contexto_geologico', '')}")
        doc.add_paragraph(f"Observações: {estacao_dict.get('observacoes', '')}")
        if structures_list:
            doc.add_heading("Estruturas", level=2)
            for s in structures_list:
                doc.add_paragraph(f"{s.get('tipo', '')} - Strike: {s.get('strike', '')}, Dip: {s.get('dip', '')}, Dip Dir: {s.get('dip_dir', '')}, Plunge: {s.get('plunge', '')}, Azimuth: {s.get('azimuth', '')}")
        if samples_list:
            doc.add_heading("Amostras", level=2)
            for s in samples_list:
                doc.add_paragraph(f"{s.get('codigo', '')} - {s.get('tipo', '')} - {s.get('finalidade', '')} - Orientada: {'Sim' if s.get('orientada') else 'Não'}")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf
    except Exception:
        html = f"""
        <h1>Estação {estacao_dict.get('ponto_id', '')}</h1>
        <p>Data: {estacao_dict.get('data', '')}</p>
        <p>Localização: {estacao_dict.get('localizacao', '')}</p>
        <p>Município: {estacao_dict.get('municipio', '')}</p>
        <p>Coordenadas: {estacao_dict.get('latitude', '')}, {estacao_dict.get('longitude', '')}</p>
        <p>Altitude: {estacao_dict.get('altitude', '')}</p>
        <p>Litologia principal: {estacao_dict.get('litologia_principal', '')}</p>
        <p>Litologia secundária: {estacao_dict.get('litologia_secundaria', '')}</p>
        <p>Tipo de afloramento: {estacao_dict.get('tipo_afloramento', '')}</p>
        <p>Contexto geológico: {estacao_dict.get('contexto_geologico', '')}</p>
        <p>Observações: {estacao_dict.get('observacoes', '')}</p>
        """
        buf = io.BytesIO(html.encode("utf-8"))
        return buf

# --- DADOS DAS TABELAS DE APOIO ---

DENSIDADE_ROCHAS = [
    ("Aluvião", "1.5 - 2.0", "1.7"),
    ("Argila", "1.6 - 2.6", "2.2"),
    ("Areia", "1.6 - 2.6", "2.2"),
    ("Arenito", "2.0 - 2.7", "2.3"),
    ("Argilito", "2.3 - 2.7", "2.5"),
    ("Calcário", "2.4 - 2.8", "2.6"),
    ("Riolito", "2.4 - 2.7", "2.5"),
    ("Andesito", "2.5 - 2.8", "2.6"),
    ("Granito", "2.5 - 2.8", "2.6"),
    ("Granodiorito", "2.6 - 2.8", "2.7"),
    ("Diabásio", "2.8 - 3.1", "2.9"),
    ("Basalto", "2.7 - 3.2", "2.9"),
    ("Gabro", "2.8 - 3.1", "2.9"),
    ("Peridotito", "2.8 - 3.4", "3.1"),
    ("Piroxenito", "2.9 - 3.2", "3.0"),
    ("Quartzito", "2.6 - 2.8", "2.7"),
    ("Xisto", "2.5 - 2.8", "2.6"),
    ("Granulito", "2.6 - 2.9", "2.7"),
    ("Filito", "2.7 - 2.8", "2.7"),
    ("Mármore", "2.5 - 2.8", "2.6"),
    ("Ardósia", "2.7 - 2.8", "2.7"),
    ("Gnaisse", "2.6 - 2.9", "2.7"),
    ("Anfibolito", "2.9 - 3.2", "3.0"),
    ("Eclogito", "3.3 - 3.5", "3.4"),
]

VELOCIDADE_ONDAS_P = [
    ("Areia seca", "0.2 - 1.0"),
    ("Areia saturada", "1.2 - 1.8"),
    ("Argila", "1.0 - 2.5"),
    ("Till glacial", "1.5 - 2.7"),
    ("Permafroste", "2.5 - 3.7"),
    ("Arenitos", "2.0 - 4.0"),
    ("Arenito Terciário", "1.4 - 2.3"),
    ("Arenito Pennant", "3.7 - 4.1"),
    ("Quartzito Cambriano", "5.5 - 5.7"),
    ("Calcários", "2.5 - 6.1"),
    ("Greda Cretácea", "2.1 - 2.8"),
    ("Oólitos Jurássicos", "2.8 - 3.7"),
    ("Calcário Carbonífero", "3.5 - 5.6"),
    ("Dolomitos", "2.5 - 6.9"),
    ("Sal", "4.5 - 5.5"),
    ("Anidrita", "4.5 - 6.2"),
    ("Gipso", "2.0 - 3.5"),
    ("Granito", "4.8 - 5.6"),
    ("Gabro", "5.5 - 6.5"),
    ("Rochas ultramáficas", "6.5 - 8.2"),
    ("Serpentinito", "5.3 - 6.6"),
    ("Ar", "0.33"),
    ("Água", "1.43 - 1.66"),
    ("Gelo", "3.4 - 3.7"),
    ("Petróleo", "1.3 - 1.4"),
    ("Gás", "0.4 - 0.6"),
    ("Aço", "5.9 - 6.1"),
    ("Ferro", "5.9 - 6.1"),
    ("Alumínio", "6.3 - 6.4"),
    ("Concreto", "3.6 - 4.3"),
]

SUSCETIBILIDADE_MAGNETICA = [
    ("Ar", "0"),
    ("Quartzo", "0"),
    ("Rocha de sal", "-0.1"),
    ("Calcita", "-0.1"),
    ("Esfalerita", "-0.1"),
    ("Pirita", "1.0"),
    ("Hematita", "0.5 - 2.5"),
    ("Ilmenita", "1.0 - 2.5"),
    ("Magnetita", "1.2 - 19.2"),
    ("Calcário", "0.0 - 0.3"),
    ("Arenito", "0.0 - 2.0"),
    ("Folhelho", "0.0 - 0.9"),
    ("Xisto", "0.0 - 0.8"),
    ("Gnaisse", "0.0 - 1.5"),
    ("Ardósia", "0.0 - 0.7"),
    ("Granito", "0.0 - 1.0"),
    ("Gabro", "0.5 - 4.0"),
    ("Basalto", "0.2 - 6.0"),
    ("Peridotito", "0.5 - 6.0"),
]

# --- INICIALIZAÇÃO ---
init_db()
init_license()

if "pagina" not in st.session_state:
    st.session_state.pagina = "🆕 Nova Estação"
if "edit_id" not in st.session_state:
    st.session_state.edit_id = None

# --- SIDEBAR ---
st.sidebar.image("https://cdn-icons-png.flaticon.com/512/3063/3063176.png", width=80)
st.sidebar.title("AfloraGeo 🌋")
st.sidebar.markdown("Caderneta de campo geológica brasileira")

licenca = get_license()
station_count = get_station_count()
limit = licenca["stations_limit"] if licenca else 30

with st.sidebar.container(border=True):
    st.markdown("### 💎 Status da Licença")
    if licenca and licenca["is_premium"]:
        st.success("💎 Premium ativo")
        st.write(f"Plano: {licenca['plan_type']}")
        st.write(f"Email: {licenca['user_email']}")
        st.write(f"Expira em: {licenca['expires_at']}")
    else:
        st.write(f"Free - {station_count}/{limit} estações")
        ratio = min(station_count / limit, 1.0) if limit > 0 else 0
        if station_count <= 20:
            color = "green"
        elif station_count <= 28:
            color = "yellow"
        else:
            color = "red"
        st.progress(ratio, text=f"{station_count}/{limit} ({color})")

pagina = st.sidebar.radio(
    "Navegação",
    ["🆕 Nova Estação", "📋 Lista de Estações", "🗺️ Mapa", "📊 Tabelas de Apoio", "📤 Exportar", "💎 Premium", "⚙️ Configurações"],
    index=["🆕 Nova Estação", "📋 Lista de Estações", "🗺️ Mapa", "📊 Tabelas de Apoio", "📤 Exportar", "💎 Premium", "⚙️ Configurações"].index(st.session_state.pagina),
)
if pagina != st.session_state.pagina:
    st.session_state.pagina = pagina
    if pagina != "🆕 Nova Estação":
        st.session_state.edit_id = None
    st.rerun()

st.sidebar.markdown("---")
st.sidebar.markdown("v1.0 - AfloraGeo")

# --- PÁGINA: NOVA ESTAÇÃO ---
if st.session_state.pagina == "🆕 Nova Estação":
    edit_id = st.session_state.edit_id
    editing = edit_id is not None
    estacao = None
    structures_existing = []
    samples_existing = []
    if editing:
        estacao = get_station_by_id(edit_id)
        structures_existing = get_structures_by_station(edit_id)
        samples_existing = get_samples_by_station(edit_id)

    if editing and estacao:
        st.title(f"✏️ Editar Estação {estacao.get('ponto_id', '')}")
    else:
        st.title("🆕 Nova Estação")

    if not editing and not can_add_station():
        st.error("Limite de estações atingido no plano Free. Assine o Premium para adicionar mais estações.")
        if st.button("💎 Assinar Premium", type="primary"):
            st.session_state.pagina = "💎 Premium"
            st.rerun()
        st.stop()

    with st.form("form_estacao"):
        col1, col2 = st.columns(2)
        with col1:
            ponto_id = st.text_input("Nº do ponto", value=estacao.get("ponto_id") if estacao else "", disabled=True, help="Gerado automaticamente")
        with col2:
            data = st.date_input("Data", value=datetime.datetime.strptime(estacao.get("data"), "%Y-%m-%d").date() if estacao and estacao.get("data") else datetime.date.today())

        with st.expander("📍 Coordenadas", expanded=True):
            coord_tipo = st.radio("Sistema de coordenadas", ["UTM", "DD (Latitude/Longitude)"], index=0 if estacao and estacao.get("utm_east") else 1)
            if coord_tipo == "UTM":
                c1, c2, c3, c4 = st.columns(4)
                with c1:
                    utm_zone = st.text_input("Zona UTM", value=estacao.get("utm_zone") if estacao else "23K")
                with c2:
                    hemisferio = st.selectbox("Hemisfério", ["S", "N"], index=0 if estacao and estacao.get("hemisferio") == "S" else 1)
                with c3:
                    utm_east = st.number_input("Este (m)", value=float(estacao.get("utm_east")) if estacao and estacao.get("utm_east") else 0.0, format="%.2f")
                with c4:
                    utm_north = st.number_input("Norte (m)", value=float(estacao.get("utm_north")) if estacao and estacao.get("utm_north") else 0.0, format="%.2f")
                lat, lon = utm_to_dd(utm_zone, utm_east, utm_north, hemisferio)
                if lat is not None and lon is not None:
                    st.success(f"Latitude: {lat:.6f}, Longitude: {lon:.6f}")
                else:
                    lat, lon = None, None
                    st.warning("Não foi possível converter UTM para DD. Verifique zona, este e norte.")
            else:
                c1, c2 = st.columns(2)
                with c1:
                    lat = st.number_input("Latitude", value=float(estacao.get("latitude")) if estacao and estacao.get("latitude") else 0.0, format="%.6f")
                with c2:
                    lon = st.number_input("Longitude", value=float(estacao.get("longitude")) if estacao and estacao.get("longitude") else 0.0, format="%.6f")
                zone, hemisferio, utm_east, utm_north = dd_to_utm(lat, lon)
                if zone:
                    st.success(f"Zona UTM: {zone} ({hemisferio}), Este: {utm_east:.2f}, Norte: {utm_north:.2f}")
                else:
                    utm_zone, hemisferio, utm_east, utm_north = "", "S", 0.0, 0.0

        altitude = st.number_input("Altitude (m)", value=float(estacao.get("altitude")) if estacao and estacao.get("altitude") else 0.0, format="%.2f")
        localizacao = st.text_input("Localização", value=estacao.get("localizacao") if estacao else "")
        municipio = st.text_input("Município", value=estacao.get("municipio") if estacao else "")
        contexto_geologico = st.text_area("Contexto geológico", value=estacao.get("contexto_geologico") if estacao else "")

        c1, c2, c3, c4 = st.columns(4)
        with c1:
            tipo_afloramento = st.selectbox(
                "Tipo de afloramento",
                ["Corte de estrada", "Lajeado", "Pedreira", "Rio", "Trincheira", "Galeria", "Outro"],
                index=["Corte de estrada", "Lajeado", "Pedreira", "Rio", "Trincheira", "Galeria", "Outro"].index(estacao.get("tipo_afloramento")) if estacao and estacao.get("tipo_afloramento") in ["Corte de estrada", "Lajeado", "Pedreira", "Rio", "Trincheira", "Galeria", "Outro"] else 0,
            )
        with c2:
            dimensoes = st.text_input("Dimensões (ex: 10 x 5 m)", value=estacao.get("dimensoes") if estacao else "")
        with c3:
            orientacao_afloramento = st.text_input("Orientação do afloramento", value=estacao.get("orientacao_afloramento") if estacao else "")
        with c4:
            acesso_opts = ["caminhamento", "carro", "barco", "helicoptero"]
            acesso_defaults = [a.strip() for a in (estacao.get("acesso") or "").split(",") if a.strip() in acesso_opts] if estacao else []
            acesso = ",".join(st.multiselect("Acesso", acesso_opts, default=acesso_defaults))

        litologias = [
            "Arenito", "Folhelho", "Calcário", "Diamictito", "Basalto", "Granito", "Gnaisse", "Xisto", "Quartzito",
            "Mármore", "Argilito", "Siltito", "Conglomerado", "Riolito", "Gabro", "Anfibolito", "Outro"
        ]
        c1, c2 = st.columns(2)
        with c1:
            litologia_principal = st.selectbox(
                "Litologia principal",
                litologias,
                index=litologias.index(estacao.get("litologia_principal")) if estacao and estacao.get("litologia_principal") in litologias else 0,
            )
            textura_principal = st.text_input("Textura litologia principal", value=estacao.get("litologia_principal") if estacao and estacao.get("litologia_principal") and estacao.get("litologia_principal") not in litologias else "")
        with c2:
            litologia_secundaria = st.selectbox(
                "Litologia secundária",
                ["Nenhuma"] + litologias,
                index=(["Nenhuma"] + litologias).index(estacao.get("litologia_secundaria")) if estacao and estacao.get("litologia_secundaria") in (["Nenhuma"] + litologias) else 0,
            )
            textura_secundaria = st.text_input("Textura litologia secundária", value=estacao.get("litologia_secundaria") if estacao and estacao.get("litologia_secundaria") and estacao.get("litologia_secundaria") not in litologias else "")

        c1, c2, c3 = st.columns(3)
        with c1:
            granulometria = st.selectbox(
                "Granulometria",
                ["Argila", "Silte", "Areia fina", "Areia média", "Areia grossa", "Areia muito grossa", "Seixo", "Calhau", "Matacão"],
                index=["Argila", "Silte", "Areia fina", "Areia média", "Areia grossa", "Areia muito grossa", "Seixo", "Calhau", "Matacão"].index(estacao.get("granulometria")) if estacao and estacao.get("granulometria") in ["Argila", "Silte", "Areia fina", "Areia média", "Areia grossa", "Areia muito grossa", "Seixo", "Calhau", "Matacão"] else 0,
            )
        with c2:
            cor = st.text_input("Cor", value=estacao.get("cor") if estacao else "")
        with c3:
            intemperismo = st.selectbox(
                "Intemperismo ISRM",
                ["FR", "FRa", "FRr", "FM", "FMr", "FD", "FDr"],
                index=["FR", "FRa", "FRr", "FM", "FMr", "FD", "FDr"].index(estacao.get("intemperismo")) if estacao and estacao.get("intemperismo") in ["FR", "FRa", "FRr", "FM", "FMr", "FD", "FDr"] else 0,
            )

        with st.expander("📐 Estruturas e Atitudes"):
            st.markdown("Adicione as estruturas da estação abaixo.")
            structures = []
            num_structures = max(len(structures_existing), 1)
            for i in range(num_structures):
                existing = structures_existing[i] if i < len(structures_existing) else {}
                with st.container(border=True):
                    c1, c2, c3, c4, c5, c6, c7 = st.columns([2, 1, 1, 1, 1, 1, 2])
                    with c1:
                        s_tipo = st.selectbox(
                            "Tipo",
                            ["Acamamento", "Foliação", "Fratura", "Falha", "Dobra", "Lineação", "Paleocorrente", "Veio", "Dique", "Xistosidade", "Clivagem", "Eixo de dobra", "Superfície axial", "Outro"],
                            index=["Acamamento", "Foliação", "Fratura", "Falha", "Dobra", "Lineação", "Paleocorrente", "Veio", "Dique", "Xistosidade", "Clivagem", "Eixo de dobra", "Superfície axial", "Outro"].index(existing.get("tipo")) if existing.get("tipo") in ["Acamamento", "Foliação", "Fratura", "Falha", "Dobra", "Lineação", "Paleocorrente", "Veio", "Dique", "Xistosidade", "Clivagem", "Eixo de dobra", "Superfície axial", "Outro"] else 0,
                            key=f"s_tipo_{i}",
                        )
                    with c2:
                        s_strike = st.text_input("Strike", value=existing.get("strike") or "", key=f"s_strike_{i}")
                    with c3:
                        s_dip = st.number_input("Dip", value=float(existing.get("dip")) if existing.get("dip") else 0.0, key=f"s_dip_{i}")
                    with c4:
                        s_dip_dir = st.number_input("Dip Dir", value=float(existing.get("dip_dir")) if existing.get("dip_dir") else 0.0, key=f"s_dip_dir_{i}")
                    with c5:
                        s_plunge = st.number_input("Plunge", value=float(existing.get("plunge")) if existing.get("plunge") else 0.0, key=f"s_plunge_{i}")
                    with c6:
                        s_azimuth = st.number_input("Azimuth", value=float(existing.get("azimuth")) if existing.get("azimuth") else 0.0, key=f"s_azimuth_{i}")
                    with c7:
                        s_obs = st.text_input("Observações", value=existing.get("observacoes") or "", key=f"s_obs_{i}")
                    structures.append({
                        "tipo": s_tipo, "strike": s_strike, "dip": s_dip, "dip_dir": s_dip_dir,
                        "plunge": s_plunge, "azimuth": s_azimuth, "observacoes": s_obs
                    })

        with st.expander("💎 Amostras"):
            st.markdown("Adicione as amostras coletadas.")
            samples_list = []
            num_samples = max(len(samples_existing), 1)
            for i in range(num_samples):
                existing = samples_existing[i] if i < len(samples_existing) else {}
                with st.container(border=True):
                    c1, c2, c3, c4, c5 = st.columns([2, 1, 1, 1, 2])
                    with c1:
                        sa_codigo = st.text_input("Código", value=existing.get("codigo") or "", key=f"sa_codigo_{i}")
                    with c2:
                        sa_tipo = st.selectbox(
                            "Tipo",
                            ["Manual", "Testemunho", "Calha"],
                            index=["Manual", "Testemunho", "Calha"].index(existing.get("tipo")) if existing.get("tipo") in ["Manual", "Testemunho", "Calha"] else 0,
                            key=f"sa_tipo_{i}",
                        )
                    with c3:
                        sa_finalidade = st.selectbox(
                            "Finalidade",
                            ["Petrografia", "Geoquímica", "Geocronologia", "Granulometria", "Outro"],
                            index=["Petrografia", "Geoquímica", "Geocronologia", "Granulometria", "Outro"].index(existing.get("finalidade")) if existing.get("finalidade") in ["Petrografia", "Geoquímica", "Geocronologia", "Granulometria", "Outro"] else 0,
                            key=f"sa_finalidade_{i}",
                        )
                    with c4:
                        sa_orientada = st.checkbox("Orientada", value=bool(existing.get("orientada")), key=f"sa_orientada_{i}")
                    with c5:
                        sa_obs = st.text_input("Observações", value=existing.get("observacoes") or "", key=f"sa_obs_{i}")
                    samples_list.append({
                        "codigo": sa_codigo, "tipo": sa_tipo, "finalidade": sa_finalidade,
                        "orientada": 1 if sa_orientada else 0, "observacoes": sa_obs
                    })

        observacoes = st.text_area("Observações adicionais", value=estacao.get("observacoes") if estacao else "")

        submitted = st.form_submit_button("💾 Salvar estação", type="primary")

    if submitted:
        if not editing and not can_add_station():
            st.error("Limite atingido. Assine Premium.")
        elif coord_tipo == "UTM" and (lat is None or lon is None):
            st.error("Coordenadas UTM inválidas. Verifique zona, este e norte.")
        else:
            if litologia_principal == "Outro":
                litologia_principal = textura_principal or "Outro"
            if litologia_secundaria == "Outro":
                litologia_secundaria = textura_secundaria or "Outro"
            data_str = data.strftime("%Y-%m-%d")
            station_data = {
                "data": data_str, "utm_zone": utm_zone if coord_tipo == "UTM" else zone or "",
                "hemisferio": hemisferio if coord_tipo == "UTM" else hemisferio or "S",
                "utm_east": utm_east if coord_tipo == "UTM" else utm_east or 0.0,
                "utm_north": utm_north if coord_tipo == "UTM" else utm_north or 0.0,
                "latitude": lat, "longitude": lon, "altitude": altitude,
                "localizacao": localizacao, "municipio": municipio,
                "contexto_geologico": contexto_geologico, "tipo_afloramento": tipo_afloramento,
                "dimensoes": dimensoes, "orientacao_afloramento": orientacao_afloramento,
                "acesso": acesso, "litologia_principal": litologia_principal,
                "litologia_secundaria": litologia_secundaria, "granulometria": granulometria,
                "cor": cor, "intemperismo": intemperismo, "observacoes": observacoes,
            }
            if editing:
                ok = update_station(edit_id, station_data, structures, samples_list)
                if ok:
                    st.success("Estação atualizada com sucesso!")
                    st.session_state.edit_id = None
                    st.session_state.pagina = "📋 Lista de Estações"
                    st.rerun()
            else:
                ok = insert_station(station_data, structures, samples_list)
                if ok:
                    st.success("Estação salva com sucesso!")
                    st.session_state.pagina = "📋 Lista de Estações"
                    st.rerun()

# --- PÁGINA: LISTA DE ESTAÇÕES ---
elif st.session_state.pagina == "📋 Lista de Estações":
    st.title("📋 Lista de Estações")
    df = list_stations_df()
    if df is None or df.empty:
        st.info("Nenhuma estação cadastrada.")
    else:
        import pandas as pd
        busca = st.text_input("Buscar texto")
        c1, c2, c3 = st.columns(3)
        with c1:
            data_inicio = st.date_input("Data início", value=None)
        with c2:
            data_fim = st.date_input("Data fim", value=None)
        with c3:
            litologias_unicas = sorted(df["litologia_principal"].dropna().unique().tolist())
            filtro_litologia = st.multiselect("Litologia", litologias_unicas)

        df_filt = df.copy()
        if busca:
            mask = df_filt.astype(str).apply(lambda x: x.str.contains(busca, case=False, na=False)).any(axis=1)
            df_filt = df_filt[mask]
        if data_inicio:
            df_filt = df_filt[df_filt["data"] >= data_inicio.strftime("%Y-%m-%d")]
        if data_fim:
            df_filt = df_filt[df_filt["data"] <= data_fim.strftime("%Y-%m-%d")]
        if filtro_litologia:
            df_filt = df_filt[df_filt["litologia_principal"].isin(filtro_litologia)]

        st.dataframe(df_filt, use_container_width=True, hide_index=True)
        st.markdown(f"Total filtrado: {len(df_filt)}")

        for _, row in df_filt.iterrows():
            station_id = row["id"]
            ponto_id = row["ponto_id"]
            with st.expander(f"{ponto_id} - {row.get('data', '')} - {row.get('localizacao', '')}"):
                st.write(f"**Município:** {row.get('municipio', '')}")
                st.write(f"**Coordenadas:** {row.get('latitude', '')}, {row.get('longitude', '')}")
                st.write(f"**Litologia principal:** {row.get('litologia_principal', '')}")
                st.write(f"**Tipo de afloramento:** {row.get('tipo_afloramento', '')}")
                st.write(f"**Contexto geológico:** {row.get('contexto_geologico', '')}")
                st.write(f"**Observações:** {row.get('observacoes', '')}")
                structures = get_structures_by_station(station_id)
                if structures:
                    st.write("**Estruturas:**")
                    for s in structures:
                        st.write(f"- {s['tipo']}: strike {s['strike']}, dip {s['dip']}, dip_dir {s['dip_dir']}, plunge {s['plunge']}, azimuth {s['azimuth']}")
                samples = get_samples_by_station(station_id)
                if samples:
                    st.write("**Amostras:**")
                    for s in samples:
                        st.write(f"- {s['codigo']} ({s['tipo']}) - {s['finalidade']} - Orientada: {'Sim' if s['orientada'] else 'Não'}")
                c1, c2 = st.columns(2)
                with c1:
                    if st.button("✏️ Editar", key=f"edit_{station_id}"):
                        st.session_state.edit_id = station_id
                        st.session_state.pagina = "🆕 Nova Estação"
                        st.rerun()
                with c2:
                    with st.popover(f"🗑️ Excluir {ponto_id}"):
                        st.warning(f"Tem certeza que deseja excluir {ponto_id}?")
                        if st.button("Confirmar exclusão", key=f"del_{station_id}"):
                            delete_station(station_id)
                            st.success("Excluído!")
                            st.rerun()

# --- PÁGINA: MAPA ---
elif st.session_state.pagina == "🗺️ Mapa":
    st.title("🗺️ Mapa de Estações")
    try:
        import folium
        from streamlit_folium import st_folium
        from folium.plugins import MarkerCluster
        df = list_stations_df()
        if df is None or df.empty:
            st.info("Nenhuma estação para mostrar no mapa.")
        else:
            m = folium.Map(location=[-14.235, -51.925], zoom_start=4)
            marker_cluster = MarkerCluster().add_to(m)
            for _, row in df.iterrows():
                lat = row.get("latitude")
                lon = row.get("longitude")
                if lat and lon:
                    popup_html = f"""
                    <b>{row.get('ponto_id', '')}</b><br>
                    Litologia: {row.get('litologia_principal', '')}<br>
                    Data: {row.get('data', '')}<br>
                    Local: {row.get('localizacao', '')}
                    """
                    folium.Marker(
                        [lat, lon],
                        popup=folium.Popup(popup_html, max_width=250),
                        tooltip=row.get("ponto_id", ""),
                    ).add_to(marker_cluster)
            st_folium(m, width=1200, height=600, returned_objects=[])
            if st.button("🔄 Atualizar mapa"):
                st.rerun()
    except Exception as e:
        st.error(f"Erro ao carregar mapa: {e}")
        st.info("Certifique-se de ter folium e streamlit-folium instalados.")

# --- PÁGINA: TABELAS DE APOIO ---
elif st.session_state.pagina == "📊 Tabelas de Apoio":
    st.title("📊 Tabelas de Apoio")
    import pandas as pd
    tab1, tab2, tab3 = st.tabs(["Densidade de Rochas", "Velocidade de Ondas P", "Suscetibilidade Magnética"])
    with tab1:
        st.markdown("Fonte: Telford et al., 1990")
        df_den = pd.DataFrame(DENSIDADE_ROCHAS, columns=["Material", "Intervalo (g/cm³)", "Valor Médio (g/cm³)"])
        st.dataframe(df_den, use_container_width=True, hide_index=True)
    with tab2:
        st.markdown("Fonte: Kearey, Brooks & Hill, 2002")
        df_vp = pd.DataFrame(VELOCIDADE_ONDAS_P, columns=["Material", "Velocidade Vp (km/s)"])
        st.dataframe(df_vp, use_container_width=True, hide_index=True)
    with tab3:
        df_sus = pd.DataFrame(SUSCETIBILIDADE_MAGNETICA, columns=["Material", "Suscetibilidade (x10³ SI)"])
        st.dataframe(df_sus, use_container_width=True, hide_index=True)

# --- PÁGINA: EXPORTAR ---
elif st.session_state.pagina == "📤 Exportar":
    st.title("📤 Exportar Dados")
    import pandas as pd
    df = list_stations_df()
    if df is None or df.empty:
        st.info("Nenhuma estação para exportar.")
    else:
        structures_df = pd.DataFrame()
        samples_df = pd.DataFrame()
        try:
            conn = get_connection()
            structures_df = pd.read_sql_query("SELECT * FROM structures", conn)
            samples_df = pd.read_sql_query("SELECT * FROM samples", conn)
        except Exception:
            pass
        finally:
            if conn:
                conn.close()

        c1, c2, c3, c4 = st.columns(4)
        with c1:
            kml_data = generate_kml(df, structures_df, samples_df)
            st.download_button("📍 KML (Google Earth)", data=kml_data, file_name="aflorageo.kml", mime="application/vnd.google-earth.kml+xml")
        with c2:
            csv_data = generate_csv(df)
            st.download_button("📄 CSV", data=csv_data, file_name="aflorageo.csv", mime="text/csv")
        with c3:
            geojson_data = generate_geojson(df)
            st.download_button("🌐 GeoJSON", data=geojson_data, file_name="aflorageo.geojson", mime="application/geo+json")
        with c4:
            if is_premium():
                for _, row in df.iterrows():
                    st_id = row["id"]
                    est_dict = row.to_dict()
                    st_list = get_structures_by_station(st_id)
                    sa_list = get_samples_by_station(st_id)
                    docx_buf = generate_docx(est_dict, st_list, sa_list)
                    st.download_button(f"DOCX {row['ponto_id']}", data=docx_buf, file_name=f"{row['ponto_id']}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            else:
                st.error("Exportação DOCX disponível apenas para Premium")
                if st.button("💎 Assinar Premium", key="btn_premium_export"):
                    st.session_state.pagina = "💎 Premium"
                    st.rerun()

# --- PÁGINA: PREMIUM ---
elif st.session_state.pagina == "💎 Premium":
    st.title("💎 Premium")
    lic = get_license()
    if lic and lic["is_premium"]:
        st.success("🎉 Parabéns! Você já possui o plano Premium ativo.")
        st.write(f"**Plano:** {lic['plan_type']}")
        st.write(f"**Email:** {lic['user_email']}")
        st.write(f"**Expira em:** {lic['expires_at']}")
        if st.button("Gerenciar assinatura"):
            st.info("Gerenciamento em breve via Pix/Mercado Pago.")
    else:
        st.markdown("### Escolha seu plano")
        c1, c2, c3 = st.columns(3)
        with c1:
            with st.container(border=True):
                st.markdown("<<div class='premium-card'>Mensal<br>R$ 19,90/mês</div>", unsafe_allow_html=True)
                if st.button("Assinar Mensal", key="mensal"):
                    st.info("Pagamento em breve via Pix/Mercado Pago.")
        with c2:
            with st.container(border=True):
                st.markdown("<<div class='premium-card'>Semestral<br>R$ 99,00/semestre<br>Economia 17%</div>", unsafe_allow_html=True)
                if st.button("Assinar Semestral", key="semestral"):
                    st.info("Pagamento em breve via Pix/Mercado Pago.")
        with c3:
            with st.container(border=True):
                st.markdown("<<div class='premium-card-popular'>⭐ Mais popular<br>Anual<br>R$ 179,00/ano<br>Economia 25%</div>", unsafe_allow_html=True)
                if st.button("Assinar Anual", key="anual"):
                    st.info("Pagamento em breve via Pix/Mercado Pago.")

        st.markdown("### Comparativo Free vs Premium")
        comparativo = [
            ["Estações", "30", "Ilimitadas"],
            ["Caderneta digital", "✅", "✅"],
            ["Mapa interativo", "✅", "✅"],
            ["Tabelas de apoio", "✅", "✅"],
            ["Exportação KML/GeoJSON/CSV", "✅", "✅"],
            ["Exportação DOCX", "❌", "✅"],
            ["Estereogramas (Schmidt + Rosetas)", "❌", "✅"],
            ["Diagramas ternários", "❌", "✅"],
            ["Módulo SedLog", "❌", "✅"],
            ["Integração CPRM", "❌", "✅"],
            ["Seções geológicas", "❌", "✅"],
            ["Compartilhamento", "❌", "✅"],
        ]
        df_comp = pd.DataFrame(comparativo, columns=["Funcionalidade", "Free", "Premium"])
        st.dataframe(df_comp, use_container_width=True, hide_index=True)

        with st.expander("🔑 Ativação manual com chave"):
            chave = st.text_input("Chave de ativação", type="password")
            email_ativ = st.text_input("Email", key="email_ativacao")
            plano_ativ = st.selectbox("Plano", ["monthly", "semiannual", "annual"], key="plano_ativacao")
            if st.button("Ativar Premium"):
                if chave == "AFLORAGEO-PREMIUM-2024":
                    if activate_premium(plano_ativ, email_ativ):
                        st.success("Premium ativado com sucesso!")
                        st.rerun()
                else:
                    st.error("Chave inválida.")

# --- PÁGINA: CONFIGURAÇÕES ---
elif st.session_state.pagina == "⚙️ Configurações":
    st.title("⚙️ Configurações")
    lic = get_license()
    st.markdown("### Informações da Conta")
    st.write(f"Email: {lic.get('user_email', 'Não informado')}")
    st.markdown("### Status da Licença")
    if lic and lic["is_premium"]:
        st.success("💎 Premium")
        st.write(f"Plano: {lic['plan_type']}")
        st.write(f"Expira em: {lic['expires_at']}")
    else:
        st.info("Plano Free")
        st.write(f"Limite: {lic.get('stations_limit', 30)} estações")
    st.markdown("---")
    st.markdown("### 🚨 Zona de Perigo")
    with st.expander("Resetar todos os dados"):
        st.warning("Esta ação excluirá TODAS as estações, estruturas, amostras e fotos. Não pode ser desfeita.")
        confirmacao = st.text_input("Digite RESETAR para confirmar")
        if st.button("Resetar dados", type="secondary"):
            if confirmacao == "RESETAR":
                reset_all_data()
                reset_license()
                st.success("Todos os dados foram resetados.")
                st.session_state.edit_id = None
                st.session_state.pagina = "🆕 Nova Estação"
                st.rerun()
            else:
                st.error("Digite RESETAR corretamente.")